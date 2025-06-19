"""
Final Assembler Agent Format Generators

Specialized generators for different output formats (PDF, EPUB, DOCX, HTML, Markdown).
Each generator handles format-specific requirements and optimizations.
"""

import re
import time
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

from musequill.config.logging import get_logger
from musequill.agents.assembler.final_assembler_structures import (
    BookMetadata, FormattedDocument, DocumentFormat, ValidationStatus
)

logger = get_logger(__name__)


class BaseFormatGenerator(ABC):
    """Base class for all format generators."""
    
    def __init__(self, config):
        self.config = config
    
    @abstractmethod
    def generate(
        self, 
        content: str, 
        metadata: BookMetadata, 
        output_path: Path, 
        options: Dict[str, Any]
    ) -> Optional[FormattedDocument]:
        """Generate formatted document."""
        pass
    
    def _clean_content_for_format(self, content: str, format_type: str) -> str:
        """Clean content for specific format requirements."""
        if format_type == "html":
            return self._clean_for_html(content)
        elif format_type == "markdown":
            return self._clean_for_markdown(content)
        elif format_type == "docx":
            return self._clean_for_docx(content)
        else:
            return content
    
    def _clean_for_html(self, content: str) -> str:
        """Clean content for HTML format."""
        # Convert LaTeX-like commands to HTML
        content = re.sub(r'\\chapter\{([^}]+)\}', r'<h1>\1</h1>', content)
        content = re.sub(r'\\section\{([^}]+)\}', r'<h2>\1</h2>', content)
        content = re.sub(r'\\subsection\{([^}]+)\}', r'<h3>\1</h3>', content)
        content = re.sub(r'\\pagebreak', '<div style="page-break-after: always;"></div>', content)
        
        # Remove LaTeX environments
        content = re.sub(r'\\begin\{titlepage\}.*?\\end\{titlepage\}', '', content, flags=re.DOTALL)
        content = re.sub(r'\\thispagestyle\{[^}]+\}', '', content)
        content = re.sub(r'\\vspace\{[^}]+\}', '', content)
        content = re.sub(r'\\noindent', '', content)
        
        # Convert paragraphs
        content = re.sub(r'\n\n+', '</p><p>', content)
        content = '<p>' + content + '</p>'
        
        # Clean up empty paragraphs
        content = re.sub(r'<p>\s*</p>', '', content)
        
        return content
    
    def _clean_for_markdown(self, content: str) -> str:
        """Clean content for Markdown format."""
        # Convert LaTeX-like commands to Markdown
        content = re.sub(r'\\chapter\{([^}]+)\}', r'# \1', content)
        content = re.sub(r'\\section\{([^}]+)\}', r'## \1', content)
        content = re.sub(r'\\subsection\{([^}]+)\}', r'### \1', content)
        content = re.sub(r'\\pagebreak', '\n\n---\n\n', content)
        
        # Remove LaTeX environments and commands
        content = re.sub(r'\\begin\{titlepage\}.*?\\end\{titlepage\}', '', content, flags=re.DOTALL)
        content = re.sub(r'\\[a-zA-Z]+\{[^}]*\}', '', content)
        content = re.sub(r'\\[a-zA-Z]+', '', content)
        
        # Clean up multiple newlines
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        return content.strip()
    
    def _clean_for_docx(self, content: str) -> str:
        """Clean content for DOCX format."""
        # Similar to markdown but preserve some formatting
        content = re.sub(r'\\chapter\{([^}]+)\}', r'# \1', content)
        content = re.sub(r'\\section\{([^}]+)\}', r'## \1', content)
        content = re.sub(r'\\subsection\{([^}]+)\}', r'### \1', content)
        content = re.sub(r'\\pagebreak', '\n\n[PAGE_BREAK]\n\n', content)
        
        # Remove LaTeX environments
        content = re.sub(r'\\begin\{titlepage\}.*?\\end\{titlepage\}', '', content, flags=re.DOTALL)
        content = re.sub(r'\\[a-zA-Z]+\{[^}]*\}', '', content)
        content = re.sub(r'\\[a-zA-Z]+', '', content)
        
        return content.strip()


class MarkdownGenerator(BaseFormatGenerator):
    """Generator for Markdown format."""
    
    def generate(
        self, 
        content: str, 
        metadata: BookMetadata, 
        output_path: Path, 
        options: Dict[str, Any]
    ) -> Optional[FormattedDocument]:
        """Generate Markdown format."""
        start_time = time.time()
        
        try:
            # Clean content for Markdown
            markdown_content = self._clean_for_markdown(content)
            
            # Add YAML frontmatter if configured
            if options.get('include_yaml_frontmatter', True):
                frontmatter = self._create_yaml_frontmatter(metadata)
                markdown_content = frontmatter + '\n\n' + markdown_content
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            file_size = output_path.stat().st_size
            generation_time = time.time() - start_time
            
            return FormattedDocument(
                format_type=DocumentFormat.MARKDOWN,
                file_path=output_path,
                file_size=file_size,
                generation_time=generation_time,
                validation_status=ValidationStatus.VALID,
                word_count=len(markdown_content.split()),
                character_count=len(markdown_content)
            )
            
        except Exception as e:
            logger.error(f"Markdown generation failed: {e}")
            return None
    
    def _create_yaml_frontmatter(self, metadata: BookMetadata) -> str:
        """Create YAML frontmatter for Markdown."""
        return f"""---
title: "{metadata.title}"
author: "{metadata.author}"
genre: "{metadata.genre}"
description: "{metadata.description}"
word_count: {metadata.word_count}
chapter_count: {metadata.chapter_count}
language: "{metadata.language}"
version: "{metadata.version}"
generated: "{metadata.generation_timestamp}"
publisher: "{metadata.publisher}"
---"""


class HTMLGenerator(BaseFormatGenerator):
    """Generator for HTML format."""
    
    def generate(
        self, 
        content: str, 
        metadata: BookMetadata, 
        output_path: Path, 
        options: Dict[str, Any]
    ) -> Optional[FormattedDocument]:
        """Generate HTML format."""
        start_time = time.time()
        
        try:
            # Clean content for HTML
            html_content = self._clean_for_html(content)
            
            # Create complete HTML document
            full_html = self._create_html_document(html_content, metadata, options)
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            file_size = output_path.stat().st_size
            generation_time = time.time() - start_time
            
            return FormattedDocument(
                format_type=DocumentFormat.HTML,
                file_path=output_path,
                file_size=file_size,
                generation_time=generation_time,
                validation_status=ValidationStatus.VALID,
                word_count=len(html_content.split()),
                character_count=len(html_content)
            )
            
        except Exception as e:
            logger.error(f"HTML generation failed: {e}")
            return None
    
    def _create_html_document(self, content: str, metadata: BookMetadata, options: Dict[str, Any]) -> str:
        """Create complete HTML document."""
        css_styles = self._generate_css_styles(options) if options.get('include_css', True) else ""
        navigation = self._generate_navigation(metadata) if options.get('include_navigation', True) else ""
        
        return f"""<!DOCTYPE html>
<html lang="{metadata.language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="author" content="{metadata.author}">
    <meta name="description" content="{metadata.description}">
    <meta name="keywords" content="{', '.join(metadata.keywords)}">
    <title>{metadata.title}</title>
    {css_styles}
</head>
<body>
    <header class="book-header">
        <h1 class="book-title">{metadata.title}</h1>
        <p class="book-author">by {metadata.author}</p>
        <p class="book-meta">{metadata.genre} | {metadata.word_count:,} words | {metadata.chapter_count} chapters</p>
    </header>
    
    {navigation}
    
    <main class="book-content">
        {content}
    </main>
    
    <footer class="book-footer">
        <p>Generated by {metadata.publisher} on {metadata.generation_timestamp[:10]}</p>
    </footer>
</body>
</html>"""
    
    def _generate_css_styles(self, options: Dict[str, Any]) -> str:
        """Generate CSS styles for HTML document."""
        return """<style>
        body {
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            background-color: #ffffff;
            color: #333333;
        }
        
        .book-header {
            text-align: center;
            margin-bottom: 3rem;
            padding-bottom: 2rem;
            border-bottom: 2px solid #333;
        }
        
        .book-title {
            font-size: 2.5rem;
            font-weight: bold;
            margin-bottom: 1rem;
            color: #2c3e50;
        }
        
        .book-author {
            font-size: 1.5rem;
            font-style: italic;
            margin-bottom: 0.5rem;
        }
        
        .book-meta {
            font-size: 1rem;
            color: #666;
        }
        
        h1 {
            color: #2c3e50;
            font-size: 2rem;
            margin-top: 3rem;
            margin-bottom: 1.5rem;
            border-bottom: 1px solid #bdc3c7;
            padding-bottom: 0.5rem;
        }
        
        h2 {
            color: #34495e;
            font-size: 1.5rem;
            margin-top: 2rem;
            margin-bottom: 1rem;
        }
        
        h3 {
            color: #555;
            font-size: 1.25rem;
            margin-top: 1.5rem;
            margin-bottom: 0.75rem;
        }
        
        p {
            margin-bottom: 1rem;
            text-align: justify;
        }
        
        .book-footer {
            text-align: center;
            margin-top: 3rem;
            padding-top: 2rem;
            border-top: 1px solid #bdc3c7;
            font-size: 0.9rem;
            color: #666;
        }
        
        @media print {
            body { margin: 0; padding: 1cm; }
            .book-header { page-break-after: always; }
            h1 { page-break-before: always; }
        }
        
        @media (max-width: 600px) {
            body { padding: 1rem; }
            .book-title { font-size: 2rem; }
            h1 { font-size: 1.5rem; }
        }
    </style>"""
    
    def _generate_navigation(self, metadata: BookMetadata) -> str:
        """Generate navigation for HTML document."""
        return """<nav class="book-navigation">
        <details>
            <summary>Table of Contents</summary>
            <ul>
                <li><a href="#introduction">Introduction</a></li>
                <li><a href="#content">Main Content</a></li>
                <li><a href="#conclusion">Conclusion</a></li>
            </ul>
        </details>
    </nav>"""


class PDFGenerator(BaseFormatGenerator):
    """Generator for PDF format using ReportLab."""
    
    def generate(
        self, 
        content: str, 
        metadata: BookMetadata, 
        output_path: Path, 
        options: Dict[str, Any]
    ) -> Optional[FormattedDocument]:
        """Generate PDF format."""
        start_time = time.time()
        
        try:
            # Check if ReportLab is available
            try:
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib.pagesizes import A4, letter
            except ImportError:
                logger.warning("ReportLab not available for PDF generation")
                return None
            
            # Choose page size
            page_size = A4 if options.get('page_size', 'A4') == 'A4' else letter
            
            # Create document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=page_size,
                rightMargin=options.get('margin_inches', 1.0) * inch,
                leftMargin=options.get('margin_inches', 1.0) * inch,
                topMargin=options.get('margin_inches', 1.0) * inch,
                bottomMargin=options.get('margin_inches', 1.0) * inch
            )
            
            # Build content
            story = self._build_pdf_story(content, metadata, options)
            
            # Generate PDF
            doc.build(story)
            
            file_size = output_path.stat().st_size
            generation_time = time.time() - start_time
            estimated_pages = max(1, metadata.word_count // 300)
            
            return FormattedDocument(
                format_type=DocumentFormat.PDF,
                file_path=output_path,
                file_size=file_size,
                generation_time=generation_time,
                validation_status=ValidationStatus.VALID,
                page_count=estimated_pages,
                word_count=metadata.word_count
            )
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return None
    
    def _build_pdf_story(self, content: str, metadata: BookMetadata, options: Dict[str, Any]) -> list:
        """Build ReportLab story for PDF generation."""
        from reportlab.platypus import Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        chapter_style = ParagraphStyle(
            'CustomChapter',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            spaceBefore=30
        )
        
        # Title page
        story.append(Paragraph(metadata.title, title_style))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"by {metadata.author}", styles['Normal']))
        story.append(Spacer(1, 40))
        story.append(Paragraph(f"{metadata.genre} | {metadata.word_count:,} words", styles['Normal']))
        story.append(PageBreak())
        
        # Process content
        content_lines = content.split('\n')
        current_paragraph = []
        
        for line in content_lines:
            line = line.strip()
            
            if line.startswith('\\chapter{'):
                # Add previous paragraph
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                
                # Add chapter title
                chapter_match = re.search(r'\\chapter\{([^}]+)\}', line)
                if chapter_match:
                    story.append(Paragraph(chapter_match.group(1), chapter_style))
            
            elif line == '\\pagebreak':
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                story.append(PageBreak())
            
            elif line and not line.startswith('\\'):
                current_paragraph.append(line)
            
            elif not line and current_paragraph:
                story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                story.append(Spacer(1, 12))
                current_paragraph = []
        
        # Add final paragraph
        if current_paragraph:
            story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
        
        return story


class EPUBGenerator(BaseFormatGenerator):
    """Generator for EPUB format."""
    
    def generate(
        self, 
        content: str, 
        metadata: BookMetadata, 
        output_path: Path, 
        options: Dict[str, Any]
    ) -> Optional[FormattedDocument]:
        """Generate EPUB format."""
        start_time = time.time()
        
        try:
            # Check if ebooklib is available
            try:
                import ebooklib
                from ebooklib import epub
            except ImportError:
                logger.warning("ebooklib not available for EPUB generation")
                return None
            
            # Create EPUB book
            book = epub.EpubBook()
            
            # Set metadata
            book.set_identifier(f"musequill-{metadata.title.lower().replace(' ', '-')}")
            book.set_title(metadata.title)
            book.set_language(metadata.language)
            book.add_author(metadata.author)
            
            if metadata.description:
                book.set_cover("image.jpg", open('default_cover.jpg', 'rb').read() if Path('default_cover.jpg').exists() else b'')
            
            # Create chapters
            chapters = self._create_epub_chapters(content, metadata)
            
            # Add chapters to book
            for chapter in chapters:
                book.add_item(chapter)
            
            # Create navigation
            book.toc = tuple(chapters)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Define spine
            book.spine = ['nav'] + chapters
            
            # Write EPUB
            epub.write_epub(str(output_path), book)
            
            file_size = output_path.stat().st_size
            generation_time = time.time() - start_time
            
            return FormattedDocument(
                format_type=DocumentFormat.EPUB,
                file_path=output_path,
                file_size=file_size,
                generation_time=generation_time,
                validation_status=ValidationStatus.VALID,
                word_count=metadata.word_count
            )
            
        except Exception as e:
            logger.error(f"EPUB generation failed: {e}")
            return None
    
    def _create_epub_chapters(self, content: str, metadata: BookMetadata) -> list:
        """Create EPUB chapters from content."""
        from ebooklib import epub
        
        chapters = []
        chapter_pattern = r'\\chapter\{([^}]+)\}'
        chapter_splits = re.split(chapter_pattern, content)
        
        for i in range(1, len(chapter_splits), 2):
            if i + 1 < len(chapter_splits):
                chapter_title = chapter_splits[i]
                chapter_content = chapter_splits[i + 1]
                
                # Clean content
                chapter_content = self._clean_for_html(chapter_content)
                
                # Create EPUB chapter
                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=f'chapter_{i//2 + 1}.xhtml',
                    lang=metadata.language
                )
                
                chapter.content = f'''<html>
<head><title>{chapter_title}</title></head>
<body>
<h1>{chapter_title}</h1>
{chapter_content}
</body>
</html>'''
                
                chapters.append(chapter)
        
        return chapters


class DOCXGenerator(BaseFormatGenerator):
    """Generator for DOCX format."""
    
    def generate(
        self, 
        content: str, 
        metadata: BookMetadata, 
        output_path: Path, 
        options: Dict[str, Any]
    ) -> Optional[FormattedDocument]:
        """Generate DOCX format."""
        start_time = time.time()
        
        try:
            # Try using Pandoc first
            if self._try_pandoc_generation(content, metadata, output_path, options):
                file_size = output_path.stat().st_size
                generation_time = time.time() - start_time
                
                return FormattedDocument(
                    format_type=DocumentFormat.DOCX,
                    file_path=output_path,
                    file_size=file_size,
                    generation_time=generation_time,
                    validation_status=ValidationStatus.VALID,
                    word_count=metadata.word_count
                )
            
            # Fallback to python-docx
            return self._try_python_docx_generation(content, metadata, output_path, options, start_time)
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return None
    
    def _try_pandoc_generation(self, content: str, metadata: BookMetadata, output_path: Path, options: Dict[str, Any]) -> bool:
        """Try generating DOCX using Pandoc."""
        try:
            import pypandoc
            
            # Create temporary markdown file
            temp_md = output_path.parent / f"temp_{output_path.stem}.md"
            
            # Convert content to markdown
            markdown_content = self._clean_for_markdown(content)
            
            # Add title page
            title_page = f"""---
title: "{metadata.title}"
author: "{metadata.author}"
---

# {metadata.title}

*by {metadata.author}*

{metadata.genre} | {metadata.word_count:,} words | {metadata.chapter_count} chapters

Generated on {metadata.generation_timestamp[:10]}

---

"""
            
            final_content = title_page + markdown_content
            
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(final_content)
            
            # Convert using Pandoc
            extra_args = ['--toc'] if options.get('include_toc', True) else []
            
            pypandoc.convert_file(
                str(temp_md),
                'docx',
                outputfile=str(output_path),
                extra_args=extra_args
            )
            
            # Cleanup
            temp_md.unlink()
            
            return True
            
        except ImportError:
            logger.warning("Pandoc not available for DOCX generation")
            return False
        except Exception as e:
            logger.warning(f"Pandoc DOCX generation failed: {e}")
            return False
    
    def _try_python_docx_generation(self, content: str, metadata: BookMetadata, output_path: Path, options: Dict[str, Any], start_time: float) -> Optional[FormattedDocument]:
        """Try generating DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title page
            title = doc.add_heading(metadata.title, 0)
            title.alignment = 1  # Center
            
            author = doc.add_paragraph(f"by {metadata.author}")
            author.alignment = 1
            
            doc.add_page_break()
            
            # Process content
            content_lines = content.split('\n')
            current_paragraph = []
            
            for line in content_lines:
                line = line.strip()
                
                if line.startswith('\\chapter{'):
                    # Add previous paragraph
                    if current_paragraph:
                        doc.add_paragraph(' '.join(current_paragraph))
                        current_paragraph = []
                    
                    # Add chapter heading
                    chapter_match = re.search(r'\\chapter\{([^}]+)\}', line)
                    if chapter_match:
                        doc.add_heading(chapter_match.group(1), level=1)
                
                elif line == '\\pagebreak':
                    if current_paragraph:
                        doc.add_paragraph(' '.join(current_paragraph))
                        current_paragraph = []
                    doc.add_page_break()
                
                elif line and not line.startswith('\\'):
                    current_paragraph.append(line)
                
                elif not line and current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
            
            # Add final paragraph
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
            
            # Save document
            doc.save(str(output_path))
            
            file_size = output_path.stat().st_size
            generation_time = time.time() - start_time
            
            return FormattedDocument(
                format_type=DocumentFormat.DOCX,
                file_path=output_path,
                file_size=file_size,
                generation_time=generation_time,
                validation_status=ValidationStatus.VALID,
                word_count=metadata.word_count
            )
            
        except ImportError:
            logger.warning("python-docx not available for DOCX generation")
            return None
        except Exception as e:
            logger.error(f"python-docx generation failed: {e}")
            return None